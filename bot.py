import os
from flask import Flask, request, send_file
from telegram import Update, Bot
from telegram.ext import Dispatcher, CommandHandler, MessageHandler, Filters
from io import BytesIO
from docx import Document
from docx.shared import RGBColor

# Список имён иноагентов для зачеркивания (пример)
FOREIGN_AGENT_NAMES = ["Иван Иванов", "Мария Петрова", "John Smith", "Ресурсный центр для ЛГБТ", "Босов Катерина Евгеньевна", "Общество с ограниченной ответственностью «ЕЛКИН КАРТОН»",
    "Ресурсный центр для ЛГБТ",
    "Телеканал Дождь",
    "Anti-Corruption Foundation Inc.",
    "Deutsche Welle (Германия, Kurt-Schumacher-Strasse 3, 53113 Bonn)",
    "Hidemy.network Ltd.",
    "MRR-Fund gemeinnützige UG (haftungsbeschränkt), реализующий проект My Russian Rights",
    "The Bell",
    "BILD на русском",
    "Idel.Реалии",
    "Медиазона",
    "Алексей Навальный",
    "Алексей Венедиктов",
    "Максим Кац",
    "Илья Варламов",
    "Михаил Ходорковский",
    "Григорий Чхартишвили (Борис Акунин)",
    "Юлия Латынина",
    "Александр Невзоров",
    "Евгений Чичваркин",
    "Виктор Шендерович",
    "Ольга Романова",
    "Марк Фейгин",
    "Русская служба BBC",
    "Голос Америки",
    "Радио Свобода",
    "Медиахолдинг «Медуза»",
    "Альянс врачей",
    "ОВД-Инфо",
    "Проект",
    "Настоящее Время",
    "Репаблик",
    "МБХ Медиа",
    "Агора",
    "Команда 29",
    "Правозащита Открытки",
    "Русь сидящая",
    "Апология протеста",
    "Фонд борьбы с коррупцией",
    "Фонд защиты прав граждан",
    "Алексей Горинов",
    "Иван Жданов",
    "Леонид Волков",
    "Александр Сотник",
    "Марат Гельман",
    "Ксения Ларина",
    "Гарри Каспаров",
    "Михаил Зыгарь",
    "Иван Ключарёв",
    "Андрей Лошак",
    "Павел Лобков",
    "Лилия Чанышева",
    "Сергей Смирнов",
    "Илья Яшин",
    "Олег Кашин",
    "Сергей Пархоменко",
    "Татьяна Лазарева",
    "Сергей Гуриев",
    "Константин Эггерт",
    "Сергей Алексашенко",
    "Игорь Яковенко",
    "Александр Баунов",
    "Григорий Свердлин",
    "Валерий Печейкин",
    "Юрий Дудь",
    "Роман Доброхотов",
    "Алексей Пивоваров",
    "Канал Redakciya",
    "Пётр Верзилов",
    "Надежда Толоконникова",
    "Мария Алехина",
    "Pussy Riot",
    "Оксана Баулина",
    "Артемий Троицкий",
    "Аркадий Бабченко",
    "Эхо Москвы",
    "Новая Газета",
    "Новая Газета Европа",
    "Новая История",
    "Независимая Газета",
    "Сноб",
    "Такие дела",
    "Холод",
    "iStories (Важные истории)",
    "7x7",
    "Псковская губерния",
    "Сота",
    "Альберт Михалев",
    "Татьяна Фельгенгауэр",
    "Илья Азар",
    "Светлана Прокопьева",
    "Михаил Фишман",
    "Андрей Бильжо",
    "Дмитрий Быков",
    "Дмитрий Муратов",
    "Лев Шлосберг",
    "Елена Милашина",
    "Александр Плющев",
    "Антон Орехъ",
    "Николай Картозия",
    "Ирина Шихман",
    "Наталья Синдеева",
    "Канал Дождь",
    "Канал Навальный Live",
    "Канал Популярная политика",
    "Канал Ходорковский Live",
    "Канал Русский агент",
    "Канал Кашин",
    "Канал Шрёдингер",
    "Канал Varvara",
    "Канал SHULMAN",
    "Канал Нино Росебашвили",
    "Канал Спец",
    "Канал Bild на русском",
    "Канал Версус",
    "Канал Майор и Генерал",
    "Канал Живой Гвоздь",
    "Канал Латвия24",
    "Канал ВДудь",
    "Канал Max KATZ",
    "Канал OPEN Media",
    "Канал редакция",
    "Канал ГУЛАГу.net",
    "Канал Dissernet",
    "Канал Proekt.Media",
    "Канал Conflict Intelligence Team",
    "Канал Feminist Anti-War Resistance",
    "Канал Free Russia Foundation",
    "Канал Meduza",
    "Канал OVD-Info",
    "Канал Safehouse",
    "Канал Support for Political Prisoners. Memorial",
    "Канал Team Navalny",
    "Канал Verstka",
    "Канал Волна",
    "Канал Голос Америки",
    "Канал Idel.Реалии",
    "Канал Кавказ.Реалии",
    "Канал Крым.Реалии",
    "Канал Радио Свобода",
    "Канал Сибирь.Реалии",
    "Канал Sever.Realii",
    "Канал Южная Сеть",
    "Канал Эхо",
    "Канал Эхо Москвы",
    "Канал Эхо. Новости"]

TOKEN = os.getenv("BOT_TOKEN")
if not TOKEN:
    raise ValueError("Не найден токен BOT_TOKEN в переменных окружения")

bot = Bot(token=TOKEN)
app = Flask(__name__)
dispatcher = Dispatcher(bot, None, workers=0)

def start(update: Update, context=None):
    update.message.reply_text("Привет! Отправь мне Word-файл, я зачеркиваю в нём имена из списка.")

def handle_doc(update: Update, context=None):
    file = update.message.document.get_file()
    file_bytes = BytesIO()
    file.download(out=file_bytes)
    file_bytes.seek(0)

    doc = Document(file_bytes)

    for para in doc.paragraphs:
        original_text = para.text
        para.clear()  # Очищаем параграф, чтобы вручную вставить отформатированные куски

        i = 0
        while i < len(original_text):
            match_found = False
            for name in FOREIGN_AGENT_NAMES:
                if original_text[i:i+len(name)] == name:
                    run = para.add_run(name)
                    run.font.color.rgb = RGBColor(255, 0, 0)  # Красный цвет
                    i += len(name)
                    match_found = True
                    break
            if not match_found:
                run = para.add_run(original_text[i])
                i += 1

    output = BytesIO()
    doc.save(output)
    output.seek(0)

    update.message.reply_document(document=output, filename="highlighted.docx")

# Регистрируем обработчики
dispatcher.add_handler(CommandHandler("start", start))
dispatcher.add_handler(MessageHandler(Filters.document.mime_type("application/vnd.openxmlformats-officedocument.wordprocessingml.document"), handle_doc))

@app.route(f"/{TOKEN}", methods=["POST"])
def webhook():
    data = request.get_json(force=True)
    update = Update.de_json(data, bot)
    dispatcher.process_update(update)
    return "ok"

@app.route("/")
def index():
    return "Bot is running"

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    print(f"Starting app on port {port}")
    app.run(host="0.0.0.0", port=port)
